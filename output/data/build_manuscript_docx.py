"""
One-time setup: build manuscript.docx and supplementary_figures.docx from manuscript.md.

After running this script, edit the .docx files directly.
Run refresh_figures.py to swap in updated PNGs without touching any text.

Figures 1, 5, 6, 8 → supplementary (S1–S4).
Figures 2, 3, 4, 7, 9 → main manuscript, renumbered 1–5.
"""
import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DIR = os.path.dirname(os.path.abspath(__file__))
MD_FILE  = os.path.join(DIR, "manuscript.md")
MAIN_OUT = os.path.join(DIR, "manuscript.docx")
SUPP_OUT = os.path.join(DIR, "supplementary_figures.docx")

# ── Figure routing ─────────────────────────────────────────────────────────────
# Filename → original figure number
FILENAME_TO_FIG = {
    "protein_landscape.png":              1,
    "protein_pearson_correlation.png":    2,
    "domain_pearson_correlation.png":     3,
    "auc_by_label_method_by_round.png":   4,
    "auc_by_domain.png":                  5,
    "mse_by_domain.png":                  6,
    "figure3_label_strategy_auc.png":     7,
    "auc_by_label_method_by_domain.png":  8,
    "protein_final_auc.png":              9,
}
SUPP_FIGS    = {1, 5, 6, 8}
MAIN_FIG_MAP = {2: 1, 3: 2, 4: 3, 7: 4, 9: 5}
SUPP_FIG_MAP = {1: "S1", 5: "S2", 6: "S3", 8: "S4"}

IMG_WIDTH = 6.5


def remap_figure(m):
    n = int(m.group(1))
    if n in SUPP_FIG_MAP:
        return f"Supplementary Figure {SUPP_FIG_MAP[n]}"
    elif n in MAIN_FIG_MAP:
        return f"Figure {MAIN_FIG_MAP[n]}"
    return m.group(0)


def apply_fig_remap(text):
    return re.sub(r'\bFigure\s+(\d+)\b', remap_figure, text)


# ── Doc helpers ────────────────────────────────────────────────────────────────
def make_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
    return doc


def set_style(paragraph, size_pt, bold=False, italic=False):
    for run in paragraph.runs:
        run.font.size  = Pt(size_pt)
        run.font.bold  = bold
        run.font.italic = italic


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    sizes = {1: 14, 2: 12, 3: 11}
    set_style(p, sizes.get(level, 10), bold=True)
    return p


def add_inline_para(doc, text, size_pt=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|_[^_]+_)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = p.add_run(token[2:-2])
            run.bold = True
            run.font.size = Pt(size_pt)
        elif (token.startswith('*') and token.endswith('*')) or \
             (token.startswith('_') and token.endswith('_')):
            run = p.add_run(token[1:-1])
            run.italic = True
            run.font.size = Pt(size_pt)
        else:
            run = p.add_run(token)
            run.font.size = Pt(size_pt)
    return p


def add_figure_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = re.match(r'(\*\*(?:Supplementary )?Figure\s+\S+\.?\*\*)(.*)', text, re.DOTALL)
    if m:
        run = p.add_run(m.group(1).replace('**', ''))
        run.bold = True
        run.font.size = Pt(9)
        run2 = p.add_run(m.group(2))
        run2.font.size = Pt(9)
    else:
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.italic = True
    return p


def add_image(doc, filename):
    img_path = os.path.join(DIR, filename)
    if not os.path.exists(img_path):
        doc.add_paragraph(f"[IMAGE NOT FOUND: {filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inline = p.add_run().add_picture(img_path, width=Inches(IMG_WIDTH))
    # Tag with filename so refresh_figures.py can identify and replace this image
    docPr = inline._inline.find(qn('wp:docPr'))
    if docPr is not None:
        docPr.set('descr', filename)


def add_blockquote(doc, text):
    p = doc.add_paragraph(text, style='Quote')
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size   = Pt(10)
        run.italic      = True


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = p.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = p.add_run(token[1:-1])
            run.italic = True
        else:
            p.add_run(token)


def add_reference(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after       = Pt(3)
    tokens = re.split(r'(\*[^*]+\*)', text)
    for token in tokens:
        if token.startswith('*') and token.endswith('*'):
            run = p.add_run(token[1:-1])
            run.italic    = True
            run.font.size = Pt(9)
        else:
            run = p.add_run(token)
            run.font.size = Pt(9)


# ── Parse markdown into blocks ─────────────────────────────────────────────────
# Each block: {"type": str, "text": str, "filename": str|None, "fig_num": int|None}
# type: heading | para | image | caption | hr | blockquote | bullet | reference

def parse_md(path):
    with open(path) as f:
        lines = f.readlines()
    blocks = []
    i = 0
    in_references = False

    while i < len(lines):
        line = lines[i].rstrip('\n')

        if not line.strip():
            i += 1
            continue

        if re.match(r'^---+\s*$', line):
            blocks.append({"type": "hr"})
            i += 1
            continue

        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text  = m.group(2)
            if 'References' in text:
                in_references = True
            blocks.append({"type": "heading", "level": level, "text": text})
            i += 1
            continue

        m = re.match(r'^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$', line)
        if m:
            filename = m.group(2)
            fig_num  = FILENAME_TO_FIG.get(filename)
            blocks.append({"type": "image", "filename": filename, "fig_num": fig_num})
            i += 1
            continue

        if re.match(r'^\*\*(?:Figure|Supplementary)', line):
            blocks.append({"type": "caption", "text": line})
            i += 1
            continue

        if line.startswith('>'):
            blocks.append({"type": "blockquote", "text": line.lstrip('> ').strip()})
            i += 1
            continue

        if re.match(r'^[-*]\s+', line):
            blocks.append({"type": "bullet", "text": re.sub(r'^[-*]\s+', '', line)})
            i += 1
            continue

        # Accumulate paragraph continuation lines
        para_lines = [line]
        while i + 1 < len(lines):
            nxt = lines[i + 1].rstrip('\n')
            if (not nxt.strip()
                    or re.match(r'^#{1,4}\s', nxt)
                    or re.match(r'^\s*!\[', nxt)
                    or re.match(r'^\*\*Figure', nxt)
                    or re.match(r'^\*\*Supplementary', nxt)
                    or nxt.startswith('>')
                    or re.match(r'^[-*]\s', nxt)
                    or re.match(r'^---', nxt)
                    or (in_references and re.match(r'^\d+\.', nxt))):
                break
            i += 1
            para_lines.append(lines[i].rstrip('\n'))

        full = ' '.join(para_lines)
        btype = "reference" if (in_references and re.match(r'^\d+\.', full)) else "para"
        blocks.append({"type": btype, "text": full})
        i += 1

    return blocks


# ── Apply remapping to text blocks ────────────────────────────────────────────
def remap_blocks(blocks):
    out = []
    for b in blocks:
        b = dict(b)
        if "text" in b:
            b["text"] = apply_fig_remap(b["text"])
        out.append(b)
    return out


# ── Render a list of blocks into a doc ────────────────────────────────────────
def render_blocks(doc, blocks, skip_fig_nums=None, only_fig_nums=None):
    """
    skip_fig_nums: set of original fig numbers to skip image+caption for.
    only_fig_nums: if set, only render image+caption for these fig numbers
                   (used for supplementary doc — renders everything else too
                    unless it's an image/caption for a non-supplementary figure).
    """
    skip_fig_nums  = skip_fig_nums  or set()
    only_fig_nums  = only_fig_nums  or set()

    for b in blocks:
        btype = b["type"]

        if btype == "image":
            fig_num = b.get("fig_num")
            if only_fig_nums and fig_num not in only_fig_nums:
                continue
            if fig_num in skip_fig_nums:
                continue
            add_image(doc, b["filename"])

        elif btype == "caption":
            # Determine original fig num from the (already remapped) caption text
            # by checking the image block that preceded it — handled via peek below.
            # We use a simple heuristic: if this caption is for a figure being
            # filtered, skip it. We tag captions with fig_num during parsing.
            fig_num = b.get("fig_num")
            if only_fig_nums and fig_num not in only_fig_nums:
                continue
            if fig_num in skip_fig_nums:
                continue
            add_figure_caption(doc, b["text"])

        elif btype == "heading":
            add_heading(doc, b["text"], b["level"])

        elif btype == "hr":
            doc.add_paragraph()

        elif btype == "blockquote":
            add_blockquote(doc, b["text"])

        elif btype == "bullet":
            add_bullet(doc, b["text"])

        elif btype == "reference":
            add_reference(doc, b["text"])

        elif btype == "para":
            add_inline_para(doc, b["text"])


# ── Tag captions with their figure number ─────────────────────────────────────
def tag_captions(blocks):
    """Associate each caption block with the fig_num of the preceding image."""
    out   = list(blocks)
    last_fig = None
    for b in out:
        if b["type"] == "image":
            last_fig = b.get("fig_num")
        elif b["type"] == "caption":
            b["fig_num"] = last_fig
    return out


# ── Main ──────────────────────────────────────────────────────────────────────
raw_blocks    = parse_md(MD_FILE)
remapped      = remap_blocks(raw_blocks)
tagged        = tag_captions(remapped)

# Main manuscript: skip supplementary figures
main_doc = make_doc()
render_blocks(main_doc, tagged, skip_fig_nums=SUPP_FIGS)
main_doc.save(MAIN_OUT)
print(f"Saved {MAIN_OUT}")

# Supplementary document
supp_doc = make_doc()
supp_doc.add_heading("Supplementary Figures", level=1)
supp_doc.add_paragraph()

for orig_num, supp_label in sorted(SUPP_FIG_MAP.items()):
    for b in tagged:
        if b["type"] == "image" and b.get("fig_num") == orig_num:
            add_image(supp_doc, b["filename"])
        elif b["type"] == "caption" and b.get("fig_num") == orig_num:
            add_figure_caption(supp_doc, b["text"])

supp_doc.save(SUPP_OUT)
print(f"Saved {SUPP_OUT}")
